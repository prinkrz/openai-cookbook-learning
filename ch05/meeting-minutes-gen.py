from urllib import response

from docx import Document
from openai import OpenAI
import subprocess
import os
import whisper

client = OpenAI()

MODEL="gpt-5.4-mini"

def transcribe_audio(audio_file_path):
    with open(audio_file_path, "rb") as audio_file:
        transcription = client.audio.transcriptions.create(
            file=audio_file,
            model="whisper-1"
        )
    return transcription.text


# define a function that uses whisper on the local pc
def transcribe_audio_local_system(audio_file_path):
    # Use subprocess to call the whisper command line tool
    try:
        result = subprocess.run(
            ["whisper", audio_file_path, "--model", "base", "--output_format", "txt"],
            capture_output=True,
            text=True,
            check=True
        )
        # The output file will have the same name as the input file but with a .txt extension
        output_file_path = os.path.splitext(audio_file_path)[0] + ".txt"
        with open(output_file_path, "r") as f:
            transcription = f.read()
        return transcription
    except subprocess.CalledProcessError as e:
        print(f"Error during transcription: {e.stderr}")
        raise


def transcribe_audio_local(audio_file_path):
    model = whisper.load_model("turbo")
    result = model.transcribe(audio_file_path)
    transcription = result["text"]
    return transcription


def abstract_summary_extraction(transcription):
    response = client.chat.completions.create(
        model=MODEL,
        messages=[
            {
                "role": "system",
                "content": "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Provide avoid unnecesary details or tangential points.",
            },
            {
                "role": "user",
                "content": transcription
            },
        ],
    )

    return response.choices[0].message.content or ""

def key_point_extraction(transcription):
    response = client.chat.completions.create(
        model=MODEL,
        messages=[
            {
                "role": "system",
                "content": "You are a proficient AI with speciality in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about."
            },
            {
                "role": "user",
                "content": transcription
            },
        ],
    )
    return response.choices[0].message.content or ""



def action_item_extraction(transcription):
    response = client.chat.completions.create(
        model=MODEL,
        messages=[
            {
                "role": "system",
                "content": "You are an AI expert in analyzing conversations and extacting action items. Pleas review the text and identify any tasks, assignments, or actions that were agreed upon or mention as needing to be done. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely.",
            },
            {
                "role": "user",
                "content": transcription,
            },
        ],
    )
    return response.choices[0].message.content or ""

def sentiment_analysis(transcription):
    response = client.chat.completions.create(
        model=MODEL,
        messages=[
            {
                "role": "system",
                "content": "As an AI with expertise in language and emotional analysis, your task is to analyze the sentiment of the following text. Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases were used. Indicate whether the sentiment is generally positive , negative or neutral, and provide brief explanations for your analysis where possible.",
            },
            {
                "role": "user",
                "content": transcription
            },
        ],
    )
    return response.choices[0].message.content or ""


def meeting_minutes(transcription):
    abstract_summary = abstract_summary_extraction(transcription)
    key_points = key_point_extraction(transcription)
    action_items = action_item_extraction(transcription)
    sentiment = sentiment_analysis(transcription)

    return {
        "abstract_summary": abstract_summary,
        "key_points": key_points,
        "action_items": action_items,
        "sentiment": sentiment
    }

def save_as_docx(minutes, filename):
    doc = Document()
    for key, value in minutes.items():
        heading = " ".join(word.capitalize() for word in key.split("_"))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()

    doc.save(filename)





audio_file_path = "EarningsCall.wav"
if not audio_file_path:
    raise ValueError("Please provide a valid audio file path.")

# check if trascription already exists else transcribe the audio file
try:
    with open("transcription.txt", "r") as f:
        transcription = f.read()
except FileNotFoundError:
    transcription = transcribe_audio_local(audio_file_path)

print("-------------------Transcription---------------------")
print(transcription)


# save transcription to a text file
with open("transcription.txt", "w") as f:
    f.write(transcription)

minutes = meeting_minutes(transcription)
print(minutes)

save_as_docx(minutes, "meeting_minutes.docx")